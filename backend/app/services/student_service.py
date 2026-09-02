from typing import Any, Dict, List, Optional
from datetime import datetime
import os
import re
import shutil
from pathlib import Path

from fastapi import HTTPException, UploadFile
from sqlalchemy import or_
from sqlalchemy.orm import Session, joinedload

from app.models.application import Application
from app.models.company import Company
from app.models.interview import InterviewRequest
from app.models.job import Job
from app.models.user import User
from app.models.student import (
    StudentAchievement,
    StudentCertification,
    StudentExperience,
    StudentProfile,
    StudentSkill,
    StudentSoftSkill,
)
from app.schemas.student import (
    AchievementCreateRequest,
    AchievementUpdateRequest,
    ApplicationCreateRequest,
    CertificationCreateRequest,
    ExperienceCreateRequest,
    ExperienceUpdateRequest,
    SkillCreateRequest,
    SkillUpdateRequest,
    SoftSkillCreateRequest,
    SoftSkillUpdateRequest,
    StudentProfileUpdateRequest,
)

from app.config import settings

UPLOADS_DIR = settings.UPLOADS_DIR


class StudentService:
    @staticmethod
    def is_non_resume_content(text: str) -> bool:
        """Detect API documentation accidentally submitted as resume content."""
        normalized = text.lower()
        markers = (
            "api documentation",
            "/web/session/authenticate",
            "jsonrpc",
            "odoo returns",
            "request body",
        )
        return sum(marker in normalized for marker in markers) >= 2

    @staticmethod
    def get_student_profile(db: Session, user_id: int) -> StudentProfile:
        student = db.query(StudentProfile).filter(StudentProfile.user_id == user_id).first()
        if not student:
            user = db.query(User).filter(User.id == user_id).first()
            if user and user.role == "student":
                roll_no = user.email.split("@")[0] if "@" in user.email else user.email
                student = StudentProfile(
                    user_id=user.id,
                    email=user.email,
                    name=user.name or roll_no,
                    roll_no=roll_no,
                    department="Computer Science",
                    semester="1st"
                )
                db.add(student)
                db.commit()
                db.refresh(student)
            else:
                raise HTTPException(status_code=404, detail="Student profile not found")
        return student

    @staticmethod
    def update_student_profile(db: Session, user_id: int, update_data: StudentProfileUpdateRequest) -> StudentProfile:
        student = StudentService.get_student_profile(db, user_id)
        for field, value in update_data.model_dump(exclude_unset=True).items():
            if isinstance(value, str):
                value = value.strip() or None
            if field == "resume_text" and value and StudentService.is_non_resume_content(value):
                raise HTTPException(status_code=400, detail="API documentation cannot be saved as resume text")
            setattr(student, field, value)
        db.commit()
        db.refresh(student)
        return student

    @staticmethod
    def upload_resume(db: Session, user_id: int, file: UploadFile) -> dict:
        student = StudentService.get_student_profile(db, user_id)
        allowed_extensions = [".pdf", ".doc", ".docx"]
        ext = os.path.splitext(file.filename)[1].lower()
        if ext not in allowed_extensions:
            raise HTTPException(status_code=400, detail="Invalid file format. Use PDF, DOC, or DOCX.")

        filename = f"resume_{user_id}_{datetime.now().strftime('%Y%m%d%H%M%S')}{ext}"
        filepath = UPLOADS_DIR / "resumes" / filename
        os.makedirs(filepath.parent, exist_ok=True)
        with open(filepath, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        student.resume_path = f"/uploads/resumes/{filename}"
        extracted_text = StudentService.extract_resume_text(filepath, ext)
        student.resume_text = "" if StudentService.is_non_resume_content(extracted_text) else extracted_text
        db.commit()
        return {"resume_path": student.resume_path, "resume_text": student.resume_text or ""}

    @staticmethod
    def extract_resume_text(filepath: Path, extension: str) -> str:
        try:
            if extension == ".pdf":
                from pypdf import PdfReader
                text = "\n".join(page.extract_text() or "" for page in PdfReader(str(filepath)).pages)
            elif extension == ".docx":
                from docx import Document
                document = Document(str(filepath))
                text = "\n".join(paragraph.text for paragraph in document.paragraphs)
                text += "\n" + "\n".join(
                    " | ".join(cell.text for cell in row.cells)
                    for table in document.tables
                    for row in table.rows
                )
            else:
                return ""
        except Exception:
            return ""
        return re.sub(r"\n{3,}", "\n\n", text).strip()[:30000]

    @staticmethod
    def upload_photo(db: Session, user_id: int, file: UploadFile) -> str:
        student = StudentService.get_student_profile(db, user_id)
        allowed_extensions = [".jpg", ".jpeg", ".png"]
        ext = os.path.splitext(file.filename)[1].lower()
        if ext not in allowed_extensions:
            raise HTTPException(status_code=400, detail="Invalid file format. Use JPG, JPEG, or PNG.")

        filename = f"photo_{user_id}_{datetime.now().strftime('%Y%m%d%H%M%S')}{ext}"
        filepath = UPLOADS_DIR / "photos" / filename
        os.makedirs(filepath.parent, exist_ok=True)
        with open(filepath, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        student.photo_path = f"/uploads/photos/{filename}"
        db.commit()
        return student.photo_path

    @staticmethod
    def remove_resume(db: Session, user_id: int) -> None:
        student = StudentService.get_student_profile(db, user_id)
        if student.resume_path:
            rel_path = student.resume_path.removeprefix("/uploads/") if student.resume_path.startswith("/uploads/") else student.resume_path.lstrip("/")
            filepath = UPLOADS_DIR / rel_path
            if os.path.exists(filepath):
                os.remove(filepath)
        student.resume_path = None
        db.commit()

    @staticmethod
    def remove_photo(db: Session, user_id: int) -> None:
        student = StudentService.get_student_profile(db, user_id)
        if student.photo_path:
            rel_path = student.photo_path.removeprefix("/uploads/") if student.photo_path.startswith("/uploads/") else student.photo_path.lstrip("/")
            filepath = UPLOADS_DIR / rel_path
            if os.path.exists(filepath):
                os.remove(filepath)
        student.photo_path = None
        db.commit()

    @staticmethod
    def add_skill(db: Session, user_id: int, skill_data: SkillCreateRequest) -> StudentSkill:
        student = StudentService.get_student_profile(db, user_id)
        existing = (
            db.query(StudentSkill)
            .filter(StudentSkill.student_profile_id == student.id, StudentSkill.skill_name == skill_data.skill_name)
            .first()
        )
        if existing:
            raise HTTPException(status_code=400, detail="Skill already exists")

        new_skill = StudentSkill(
            student_profile_id=student.id,
            skill_area=skill_data.skill_area,
            skill_name=skill_data.skill_name,
            proficiency=skill_data.proficiency,
            proficiency_percent=skill_data.proficiency_percent,
        )
        db.add(new_skill)
        db.commit()
        db.refresh(new_skill)
        return new_skill

    @staticmethod
    def update_skill(db: Session, user_id: int, skill_id: int, skill_data: SkillUpdateRequest) -> StudentSkill:
        student = StudentService.get_student_profile(db, user_id)
        skill = db.query(StudentSkill).filter(
            StudentSkill.id == skill_id,
            StudentSkill.student_profile_id == student.id,
        ).first()
        if not skill:
            raise HTTPException(status_code=404, detail="Skill not found")
        for field, value in skill_data.model_dump(exclude_unset=True).items():
            setattr(skill, field, value)
        db.commit()
        db.refresh(skill)
        return skill

    @staticmethod
    def list_skills(db: Session, user_id: int) -> List[StudentSkill]:
        student = StudentService.get_student_profile(db, user_id)
        return db.query(StudentSkill).filter(StudentSkill.student_profile_id == student.id).all()

    @staticmethod
    def remove_skill(db: Session, user_id: int, skill_id: int) -> None:
        student = StudentService.get_student_profile(db, user_id)
        skill = db.query(StudentSkill).filter(StudentSkill.id == skill_id, StudentSkill.student_profile_id == student.id).first()
        if not skill:
            raise HTTPException(status_code=404, detail="Skill not found")
        db.delete(skill)
        db.commit()

    @staticmethod
    def add_experience(db: Session, user_id: int, exp_data: ExperienceCreateRequest) -> StudentExperience:
        student = StudentService.get_student_profile(db, user_id)
        new_exp = StudentExperience(
            student_profile_id=student.id,
            company_name=exp_data.company_name,
            title=exp_data.title,
            start_date=exp_data.start_date,
            end_date=exp_data.end_date,
            description=exp_data.description,
        )
        db.add(new_exp)
        db.commit()
        db.refresh(new_exp)
        return new_exp

    @staticmethod
    def list_experiences(db: Session, user_id: int) -> List[StudentExperience]:
        student = StudentService.get_student_profile(db, user_id)
        return (
            db.query(StudentExperience)
            .filter(StudentExperience.student_profile_id == student.id)
            .order_by(StudentExperience.start_date.desc())
            .all()
        )

    @staticmethod
    def update_experience(db: Session, user_id: int, exp_id: int, update_data: ExperienceUpdateRequest) -> StudentExperience:
        student = StudentService.get_student_profile(db, user_id)
        experience = db.query(StudentExperience).filter(StudentExperience.id == exp_id, StudentExperience.student_profile_id == student.id).first()
        if not experience:
            raise HTTPException(status_code=404, detail="Experience not found")
        for field, value in update_data.dict(exclude_unset=True).items():
            setattr(experience, field, value)
        db.commit()
        db.refresh(experience)
        return experience

    @staticmethod
    def delete_experience(db: Session, user_id: int, exp_id: int) -> None:
        student = StudentService.get_student_profile(db, user_id)
        experience = db.query(StudentExperience).filter(StudentExperience.id == exp_id, StudentExperience.student_profile_id == student.id).first()
        if not experience:
            raise HTTPException(status_code=404, detail="Experience not found")
        db.delete(experience)
        db.commit()

    @staticmethod
    def add_certification(db: Session, user_id: int, cert_data: CertificationCreateRequest) -> StudentCertification:
        student = StudentService.get_student_profile(db, user_id)
        new_cert = StudentCertification(
            student_profile_id=student.id,
            name=cert_data.name,
            issuer=cert_data.issuer,
            credential_url=cert_data.credential_url,
            issue_date=cert_data.issue_date,
            expiry_date=cert_data.expiry_date,
        )
        db.add(new_cert)
        db.commit()
        db.refresh(new_cert)
        return new_cert

    @staticmethod
    def list_certifications(db: Session, user_id: int) -> List[StudentCertification]:
        student = StudentService.get_student_profile(db, user_id)
        return db.query(StudentCertification).filter(StudentCertification.student_profile_id == student.id).all()

    @staticmethod
    def delete_certification(db: Session, user_id: int, cert_id: int) -> None:
        student = StudentService.get_student_profile(db, user_id)
        cert = db.query(StudentCertification).filter(StudentCertification.id == cert_id, StudentCertification.student_profile_id == student.id).first()
        if not cert:
            raise HTTPException(status_code=404, detail="Certification not found")
        db.delete(cert)
        db.commit()

    @staticmethod
    def add_soft_skill(db: Session, user_id: int, soft_skill_data: SoftSkillCreateRequest) -> StudentSoftSkill:
        student = StudentService.get_student_profile(db, user_id)
        existing = (
            db.query(StudentSoftSkill)
            .filter(StudentSoftSkill.student_profile_id == student.id, StudentSoftSkill.skill_name == soft_skill_data.skill_name)
            .first()
        )
        if existing:
            raise HTTPException(status_code=400, detail="Soft skill already exists")

        new_skill = StudentSoftSkill(
            student_profile_id=student.id,
            skill_area=soft_skill_data.skill_area,
            skill_name=soft_skill_data.skill_name,
            proficiency_percent=soft_skill_data.proficiency_percent,
        )
        db.add(new_skill)
        db.commit()
        db.refresh(new_skill)
        return new_skill

    @staticmethod
    def list_soft_skills(db: Session, user_id: int) -> List[StudentSoftSkill]:
        student = StudentService.get_student_profile(db, user_id)
        return db.query(StudentSoftSkill).filter(StudentSoftSkill.student_profile_id == student.id).all()

    @staticmethod
    def update_soft_skill(db: Session, user_id: int, soft_skill_id: int, skill_data: SoftSkillUpdateRequest) -> StudentSoftSkill:
        student = StudentService.get_student_profile(db, user_id)
        skill = db.query(StudentSoftSkill).filter(
            StudentSoftSkill.id == soft_skill_id,
            StudentSoftSkill.student_profile_id == student.id,
        ).first()
        if not skill:
            raise HTTPException(status_code=404, detail="Soft skill not found")
        for field, value in skill_data.model_dump(exclude_unset=True).items():
            setattr(skill, field, value)
        db.commit()
        db.refresh(skill)
        return skill

    @staticmethod
    def remove_soft_skill(db: Session, user_id: int, soft_skill_id: int) -> None:
        student = StudentService.get_student_profile(db, user_id)
        soft_skill = db.query(StudentSoftSkill).filter(StudentSoftSkill.id == soft_skill_id, StudentSoftSkill.student_profile_id == student.id).first()
        if not soft_skill:
            raise HTTPException(status_code=404, detail="Soft skill not found")
        db.delete(soft_skill)
        db.commit()

    @staticmethod
    def add_achievement(db: Session, user_id: int, achv_data: AchievementCreateRequest) -> StudentAchievement:
        student = StudentService.get_student_profile(db, user_id)
        new_achv = StudentAchievement(
            student_profile_id=student.id,
            title=achv_data.title,
            description=achv_data.description,
            date=achv_data.date,
        )
        db.add(new_achv)
        db.commit()
        db.refresh(new_achv)
        return new_achv

    @staticmethod
    def list_achievements(db: Session, user_id: int) -> List[StudentAchievement]:
        student = StudentService.get_student_profile(db, user_id)
        return db.query(StudentAchievement).filter(StudentAchievement.student_profile_id == student.id).order_by(StudentAchievement.date.desc()).all()

    @staticmethod
    def update_achievement(db: Session, user_id: int, achv_id: int, update_data: AchievementUpdateRequest) -> StudentAchievement:
        student = StudentService.get_student_profile(db, user_id)
        achievement = db.query(StudentAchievement).filter(StudentAchievement.id == achv_id, StudentAchievement.student_profile_id == student.id).first()
        if not achievement:
            raise HTTPException(status_code=404, detail="Achievement not found")
        for field, value in update_data.dict(exclude_unset=True).items():
            setattr(achievement, field, value)
        db.commit()
        db.refresh(achievement)
        return achievement

    @staticmethod
    def delete_achievement(db: Session, user_id: int, achv_id: int) -> None:
        student = StudentService.get_student_profile(db, user_id)
        achievement = db.query(StudentAchievement).filter(StudentAchievement.id == achv_id, StudentAchievement.student_profile_id == student.id).first()
        if not achievement:
            raise HTTPException(status_code=404, detail="Achievement not found")
        db.delete(achievement)
        db.commit()

    @staticmethod
    def list_published_jobs(
        db: Session,
        user_id: Optional[int] = None,
        title: Optional[str] = None,
        company_name: Optional[str] = None,
        location: Optional[str] = None,
        employment_type: Optional[str] = None,
        skills: Optional[List[str]] = None,
        page: int = 1,
        page_size: int = 20,
    ) -> Dict[str, Any]:
        query = db.query(Job).filter(Job.status == "published")
        if title:
            query = query.filter(Job.title.ilike(f"%{title}%"))
        if company_name:
            query = query.join(Company, Job.company_id == Company.id).filter(Company.company_name.ilike(f"%{company_name}%"))
        if location:
            query = query.filter(Job.location.ilike(f"%{location}%"))
        if employment_type:
            query = query.filter(Job.employment_type == employment_type)
        if skills:
            query = query.filter(Job.title.ilike(f"%{skills[0]}%"))

        query = query.filter(or_(Job.application_deadline.is_(None), Job.application_deadline >= datetime.now().date()))
        total = query.count()
        jobs = query.offset((page - 1) * page_size).limit(page_size).all()

        applied_job_ids = set()
        if user_id:
            student = db.query(StudentProfile).filter(StudentProfile.user_id == user_id).first()
            if student:
                app_rows = db.query(Application.job_id).filter(
                    Application.student_profile_id == student.id,
                    Application.status != "withdrawn"
                ).all()
                applied_job_ids = {r[0] for r in app_rows}

        result = []
        for job in jobs:
            result.append({
                "id": job.id,
                "company_id": job.company_id,
                "company_name": job.company.company_name if job.company else None,
                "title": job.title,
                "description": job.description,
                "requirements": job.requirements,
                "location": job.location,
                "employment_type": job.employment_type,
                "min_cgpa": float(job.min_cgpa) if job.min_cgpa is not None else None,
                "salary_min": float(job.salary_min) if job.salary_min is not None else None,
                "salary_max": float(job.salary_max) if job.salary_max is not None else None,
                "application_deadline": job.application_deadline,
                "created_at": job.created_at,
                "updated_at": job.updated_at,
                "is_applied": job.id in applied_job_ids,
            })

        return {"total": total, "items": result, "page": page, "page_size": page_size}

    @staticmethod
    def get_job_detail(db: Session, job_id: int, user_id: Optional[int] = None) -> Dict[str, Any]:
        job = db.query(Job).filter(Job.id == job_id, Job.status == "published").first()
        if not job:
            raise HTTPException(status_code=404, detail="Job not found or not published")
        
        is_applied = False
        if user_id:
            student = db.query(StudentProfile).filter(StudentProfile.user_id == user_id).first()
            if student:
                existing = db.query(Application).filter(
                    Application.job_id == job_id,
                    Application.student_profile_id == student.id,
                    Application.status != "withdrawn"
                ).first()
                if existing:
                    is_applied = True

        return {
            "id": job.id,
            "company_id": job.company_id,
            "company_name": job.company.company_name if job.company else "",
            "title": job.title,
            "description": job.description,
            "requirements": job.requirements,
            "location": job.location,
            "employment_type": job.employment_type,
            "min_cgpa": float(job.min_cgpa) if job.min_cgpa is not None else None,
            "salary_min": float(job.salary_min) if job.salary_min is not None else None,
            "salary_max": float(job.salary_max) if job.salary_max is not None else None,
            "application_deadline": job.application_deadline,
            "created_at": job.created_at,
            "updated_at": job.updated_at,
            "is_applied": is_applied,
        }

    @staticmethod
    def apply_to_job(db: Session, user_id: int, job_id: int, application_data: ApplicationCreateRequest) -> Dict[str, Any]:
        student = StudentService.get_student_profile(db, user_id)
        job = (
            db.query(Job)
            .options(joinedload(Job.company))
            .filter(Job.id == job_id, Job.status == "published")
            .first()
        )
        if not job:
            raise HTTPException(status_code=404, detail="Job not found or not published")

        existing = db.query(Application).filter(Application.job_id == job_id, Application.student_profile_id == student.id).first()
        if existing:
            if existing.status != "withdrawn":
                raise HTTPException(status_code=400, detail="You have already applied for this job.")
            else:
                existing.status = "applied"
                existing.cover_letter = application_data.cover_letter
                existing.resume_path = student.resume_path
                existing.updated_at = datetime.now()
                db.commit()
                db.refresh(existing)
                application = existing
        else:
            if not student.resume_path:
                raise HTTPException(status_code=400, detail="Please upload a resume before applying")

            application = Application(
                job_id=job_id,
                student_profile_id=student.id,
                cover_letter=application_data.cover_letter,
                resume_path=student.resume_path,
                status="applied",
            )
            db.add(application)
            db.commit()
            db.refresh(application)

        db.refresh(job)
        company_name = job.company.company_name if (job and job.company) else ""
        job_title = job.title if job else ""

        return {
            "id": application.id,
            "job_id": application.job_id,
            "job_title": job_title,
            "company_name": company_name,
            "cover_letter": application.cover_letter,
            "resume_path": application.resume_path,
            "status": application.status,
            "applied_at": application.created_at.isoformat() if application.created_at else None,
            "updated_at": application.updated_at.isoformat() if application.updated_at else None,
        }

    @staticmethod
    def list_applications(db: Session, user_id: int, status: Optional[str] = None, page: int = 1, page_size: int = 20) -> Dict[str, Any]:
        student = StudentService.get_student_profile(db, user_id)
        query = db.query(Application).filter(Application.student_profile_id == student.id)
        if status:
            query = query.filter(Application.status == status)

        total = query.count()
        applications = query.offset((page - 1) * page_size).limit(page_size).all()

        result = []
        for app in applications:
            result.append({
                "id": app.id,
                "job_id": app.job_id,
                "job_title": app.job.title if app.job else None,
                "company_name": app.job.company.company_name if app.job and app.job.company else None,
                "cover_letter": app.cover_letter,
                "resume_path": app.resume_path,
                "status": app.status,
                "applied_at": app.created_at.isoformat() if app.created_at else None,
                "updated_at": app.updated_at.isoformat() if app.updated_at else None,
            })

        return {"total": total, "items": result, "page": page, "page_size": page_size}

    @staticmethod
    def withdraw_application(db: Session, user_id: int, application_id: int) -> None:
        student = StudentService.get_student_profile(db, user_id)
        application = db.query(Application).filter(Application.id == application_id, Application.student_profile_id == student.id).first()
        if not application:
            raise HTTPException(status_code=404, detail="Application not found")
        if application.status not in ["applied", "shortlisted"]:
            raise HTTPException(status_code=400, detail="Cannot withdraw this application")
        application.status = "withdrawn"
        db.commit()

    @staticmethod
    def list_interview_requests(db: Session, user_id: int, status: Optional[str] = None, page: int = 1, page_size: int = 20) -> Dict[str, Any]:
        student = StudentService.get_student_profile(db, user_id)
        query = db.query(InterviewRequest).filter(InterviewRequest.student_profile_id == student.id)
        if status:
            query = query.filter(InterviewRequest.status == status)

        total = query.count()
        interviews = query.offset((page - 1) * page_size).limit(page_size).all()

        result = []
        for interview in interviews:
            result.append({
                "id": interview.id,
                "company_id": interview.company_id,
                "company_name": interview.company.company_name if interview.company else None,
                "job_id": interview.job_id,
                "job_title": interview.job.title if interview.job else None,
                "message": interview.message,
                "interview_date": interview.interview_date.isoformat() if interview.interview_date else None,
                "status": interview.status,
                "created_at": interview.created_at.isoformat() if interview.created_at else None,
                "responded_at": interview.responded_at.isoformat() if interview.responded_at else None,
            })

        return {"total": total, "items": result, "page": page, "page_size": page_size}

    @staticmethod
    def accept_interview_request(db: Session, user_id: int, request_id: int) -> Dict[str, Any]:
        student = StudentService.get_student_profile(db, user_id)
        interview = (
            db.query(InterviewRequest)
            .options(joinedload(InterviewRequest.company), joinedload(InterviewRequest.job), joinedload(InterviewRequest.application))
            .filter(InterviewRequest.id == request_id, InterviewRequest.student_profile_id == student.id)
            .first()
        )
        if not interview:
            raise HTTPException(status_code=404, detail="Interview request not found")

        if interview.status != "accepted":
            if interview.status in ["declined", "cancelled"]:
                raise HTTPException(status_code=400, detail=f"Interview request is already {interview.status}")
            interview.status = "accepted"
            interview.responded_at = datetime.utcnow()
            if interview.application:
                interview.application.status = "interviewed"
            db.commit()
            db.refresh(interview)

        return {
            "id": interview.id,
            "company_id": interview.company_id,
            "company_name": interview.company.company_name if interview.company else "",
            "job_id": interview.job_id,
            "job_title": interview.job.title if interview.job else "",
            "message": interview.message,
            "interview_date": interview.interview_date.isoformat() if interview.interview_date else None,
            "status": interview.status,
            "created_at": interview.created_at.isoformat() if interview.created_at else None,
            "responded_at": interview.responded_at.isoformat() if interview.responded_at else None,
        }

    @staticmethod
    def decline_interview_request(db: Session, user_id: int, request_id: int) -> Dict[str, Any]:
        student = StudentService.get_student_profile(db, user_id)
        interview = (
            db.query(InterviewRequest)
            .options(joinedload(InterviewRequest.company), joinedload(InterviewRequest.job))
            .filter(InterviewRequest.id == request_id, InterviewRequest.student_profile_id == student.id)
            .first()
        )
        if not interview:
            raise HTTPException(status_code=404, detail="Interview request not found")

        if interview.status != "declined":
            if interview.status in ["accepted", "cancelled"]:
                raise HTTPException(status_code=400, detail=f"Interview request is already {interview.status}")
            interview.status = "declined"
            interview.responded_at = datetime.utcnow()
            db.commit()
            db.refresh(interview)

        return {
            "id": interview.id,
            "company_id": interview.company_id,
            "company_name": interview.company.company_name if interview.company else "",
            "job_id": interview.job_id,
            "job_title": interview.job.title if interview.job else "",
            "message": interview.message,
            "interview_date": interview.interview_date.isoformat() if interview.interview_date else None,
            "status": interview.status,
            "created_at": interview.created_at.isoformat() if interview.created_at else None,
            "responded_at": interview.responded_at.isoformat() if interview.responded_at else None,
        }
